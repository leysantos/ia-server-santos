"""Exportação da Especificação Técnica para DOCX (Word) — layout técnico."""

from __future__ import annotations

import io
import re
from typing import Any

from pricing.spec.tech_spec_layout import (
    HEADING_PT_SIZES,
    document_blocks_for_export,
    docx_font_name,
    merge_formatting,
)
from pricing.spec.tech_spec_models import TechSpecDocument


def _apply_body_paragraph_format(paragraph: Any, fmt: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Pt

    pf = paragraph.paragraph_format
    line_spacing = float(fmt.get("line_spacing") or 1.5)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    align = str(fmt.get("text_align") or "justify").lower()
    if align == "center":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _apply_heading_paragraph_format(paragraph: Any, fmt: dict[str, Any], level: int) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Pt

    pf = paragraph.paragraph_format
    line_spacing = float(fmt.get("line_spacing") or 1.5)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(10 if level <= 2 else 6)
    pf.space_after = Pt(4)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.keep_with_next = True


def _add_runs_from_text(
    paragraph: Any,
    text: str,
    font_name: str,
    font_size: int,
    *,
    bold: bool = False,
) -> None:
    from docx.shared import Pt

    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        run.font.name = font_name
        run.font.size = Pt(font_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold


def _configure_document_styles(document: Any, fmt: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Pt

    font_name = docx_font_name(str(fmt.get("font_family") or "Arial"))
    font_size = int(fmt.get("font_size") or 11)
    line_spacing = float(fmt.get("line_spacing") or 1.5)

    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)
    npf = normal.paragraph_format
    npf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    npf.line_spacing = line_spacing
    npf.space_before = Pt(0)
    npf.space_after = Pt(3)
    npf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, before, after in (
        ("Heading 1", HEADING_PT_SIZES[1], 12, 6),
        ("Heading 2", HEADING_PT_SIZES[2], 10, 4),
        ("Heading 3", HEADING_PT_SIZES[3], 8, 3),
        ("Heading 4", HEADING_PT_SIZES[4], 6, 3),
    ):
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True
        spf = style.paragraph_format
        spf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        spf.line_spacing = 1.15
        spf.space_before = Pt(before)
        spf.space_after = Pt(after)
        spf.keep_with_next = True


def _add_header(section: Any, fmt: dict[str, Any], doc_title: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    header_text = str(fmt.get("header_text") or fmt.get("document_title") or doc_title or "").strip()
    if not header_text:
        return

    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(header_text)
    run.font.name = docx_font_name(str(fmt.get("font_family") or "Arial"))
    run.font.size = Pt(9)
    run.bold = True


def _add_page_number_footer(section: Any, fmt: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.clear()
    position = str(fmt.get("page_number_position") or "left").lower()
    if position == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif position == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_name = docx_font_name(str(fmt.get("font_family") or "Arial"))

    run = paragraph.add_run("Página ")
    run.font.name = font_name
    run.font.size = Pt(9)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def export_tech_spec_docx(doc: TechSpecDocument) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise ImportError("python-docx necessário: pip install python-docx") from exc

    fmt = merge_formatting(doc.formatting)
    font_name = docx_font_name(str(fmt.get("font_family") or "Arial"))
    font_size = int(fmt.get("font_size") or 11)

    document = Document()
    _configure_document_styles(document, fmt)

    for section in document.sections:
        section.top_margin = Cm(float(fmt.get("margin_top_cm") or 3.0))
        section.bottom_margin = Cm(float(fmt.get("margin_bottom_cm") or 2.0))
        section.left_margin = Cm(float(fmt.get("margin_left_cm") or 3.0))
        section.right_margin = Cm(float(fmt.get("margin_right_cm") or 2.0))
        section.header_distance = Cm(float(fmt.get("header_height_cm") or 1.2))
        section.footer_distance = Cm(float(fmt.get("footer_height_cm") or 1.0))
        display_title = str(fmt.get("document_title") or doc.title or "")
        _add_header(section, fmt, display_title)
        if fmt.get("page_numbers", True):
            _add_page_number_footer(section, fmt)

    logo_text = fmt.get("logo_text")
    if logo_text:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"[LOGO: {logo_text}]")
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.italic = True

    display_title = fmt.get("document_title") or doc.title
    if display_title:
        title_p = document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(12)
        title_p.paragraph_format.line_spacing = 1.15
        run = title_p.add_run(str(display_title))
        run.font.name = font_name
        run.font.size = Pt(14)
        run.bold = True

    blocks = document_blocks_for_export(doc.html_content, doc.markdown)

    for block in blocks:
        btype = block.get("type")
        text = str(block.get("text") or "")
        if not text:
            continue
        if btype == "heading":
            level = min(max(int(block.get("level") or 1), 1), 4)
            p = document.add_paragraph()
            _apply_heading_paragraph_format(p, fmt, level)
            size = HEADING_PT_SIZES.get(level, font_size)
            _add_runs_from_text(p, text, font_name, size, bold=True)
        elif btype == "list_item":
            p = document.add_paragraph(style="List Bullet")
            _apply_body_paragraph_format(p, fmt)
            _add_runs_from_text(p, text, font_name, font_size)
        else:
            p = document.add_paragraph()
            _apply_body_paragraph_format(p, fmt)
            _add_runs_from_text(p, text, font_name, font_size)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
