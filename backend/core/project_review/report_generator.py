"""Geração de relatórios DOCX/PDF (Módulos M, N, O, P, Q, R, V)."""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def build_review_report_docx(
    *,
    project_name: str,
    review: dict[str, Any],
    scores: dict[str, float],
    nonconformities: list[dict[str, Any]],
    analysis: dict[str, Any],
) -> bytes:
    """Relatório de Revisão Técnica (Módulo M)."""
    doc = Document()
    _add_header(doc, f"Relatório de Revisão Técnica — {project_name}")

    sections = [
        ("1. Dados Gerais", _dados_gerais(project_name, review)),
        ("2. Escopo", analysis.get("resumo") or "Revisão multidisciplinar automatizada."),
        ("3. Documentos Recebidos", _docs_recebidos(analysis)),
        ("4. Metodologia", "Ingestão automática, OCR, visão multimodal, RAG normativo e ProjectReviewAgent."),
        ("5. Análises Executadas", _analises_executadas(analysis)),
        ("6. Não Conformidades", _nc_summary(nonconformities)),
        ("7. Riscos", _list_items(analysis.get("riscos"))),
        ("8. Recomendações", _list_items(analysis.get("recomendacoes"), key="texto")),
        ("9. Conclusão", _conclusao(scores)),
    ]

    for title, body in sections:
        doc.add_heading(title, level=1)
        if isinstance(body, list):
            for line in body:
                doc.add_paragraph(str(line))
        else:
            doc.add_paragraph(str(body))

    doc.add_heading("Indicadores de Conformidade", level=1)
    for key, val in scores.items():
        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {val}/100")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_nc_report_docx(
    *,
    project_name: str,
    nonconformities: list[dict[str, Any]],
) -> bytes:
    """Relatório detalhado de NCs (Módulo N)."""
    doc = Document()
    _add_header(doc, f"Relatório de Não Conformidades — {project_name}")

    for nc in nonconformities:
        doc.add_heading(f"{nc.get('codigo')} — {nc.get('categoria', '').upper()}", level=2)
        doc.add_paragraph(f"Criticidade: {nc.get('criticidade', '—')}")
        doc.add_paragraph(f"Descrição: {nc.get('descricao', '—')}")
        doc.add_paragraph(f"Evidência: {nc.get('evidencia') or '—'}")
        doc.add_paragraph(f"Norma: {nc.get('norma') or '—'}")
        doc.add_paragraph(f"Impacto: {nc.get('impacto') or '—'}")
        doc.add_paragraph(f"Recomendação: {nc.get('recomendacao') or '—'}")
        doc.add_paragraph(f"Status: {nc.get('status', 'aberta')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_technical_opinion_docx(
    *,
    project_name: str,
    analysis: dict[str, Any],
    responsible: str = "Responsável Técnico — IA Server Santos",
) -> bytes:
    """Parecer técnico (Módulo P)."""
    doc = Document()
    _add_header(doc, f"Parecer Técnico — {project_name}")

    for title, content in [
        ("Objeto", f"Revisão técnica do empreendimento {project_name}"),
        ("Fundamentação", _normas_text(analysis)),
        ("Análise", analysis.get("resumo") or "Análise conforme documentação recebida."),
        ("Conclusão", _conclusao_parecer(analysis)),
        ("Responsável Técnico", responsible),
    ]:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_memorial_docx(
    *,
    project_name: str,
    discipline: str,
    twin_payload: dict[str, Any],
) -> bytes:
    """Memorial descritivo automático (Módulo Q)."""
    doc = Document()
    _add_header(doc, f"Memorial Descritivo {discipline.title()} — {project_name}")

    block = (twin_payload.get("payload") or twin_payload).get(discipline) or {}
    doc.add_paragraph(f"Projeto: {project_name}")
    doc.add_paragraph(f"Disciplina: {discipline}")
    for doc_item in block.get("documentos") or []:
        doc.add_paragraph(f"• Documento: {doc_item.get('filename')}")
    for el in block.get("elementos") or []:
        doc.add_paragraph(f"• Elemento: {el}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_tdr_docx(*, project_name: str, scope: str = "") -> bytes:
    """Termo de Referência (Módulo R)."""
    doc = Document()
    _add_header(doc, f"Termo de Referência — {project_name}")

    sections = [
        ("Objeto", f"Contratação de serviços de engenharia para {project_name}"),
        ("Justificativa", "Necessidade de adequação técnica e conformidade normativa."),
        ("Escopo", scope or "Elaboração/revisão de projetos compatibilizados e documentação técnica."),
        ("Especificações", "Conforme normas ABNT, legislação aplicável e memorial de cálculo."),
        ("Medição", "Por planilha analítica e boletim de medição."),
        ("Pagamento", "Conforme cronograma físico-financeiro."),
        ("Fiscalização", "Acompanhamento técnico pelo contratante."),
        ("Recebimento", "Mediante relatório de conformidade e ART/RRT."),
        ("Cronograma", "A definir conforme complexidade do empreendimento."),
    ]
    for title, text in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_docx_to_path(content: bytes, path: Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(content)
    return path


def _add_header(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph(f"Emitido em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")


def _dados_gerais(project_name: str, review: dict[str, Any]) -> str:
    return (
        f"Projeto: {project_name}\n"
        f"Revisão: v{review.get('version', 1)}\n"
        f"Status: {review.get('status', '—')}"
    )


def _docs_recebidos(analysis: dict[str, Any]) -> str:
    missing = analysis.get("documentos_faltantes") or []
    if missing:
        return "Documentos faltantes: " + ", ".join(str(m) for m in missing)
    return "Documentação recebida conforme escopo analisado."


def _analises_executadas(analysis: dict[str, Any]) -> str:
    parts = []
    for key in ("inconsistencias", "omissoes", "divergencias", "conflitos"):
        items = analysis.get(key) or []
        parts.append(f"{key}: {len(items)} ocorrência(s)")
    return "\n".join(parts)


def _nc_summary(ncs: list[dict[str, Any]]) -> str:
    if not ncs:
        return "Nenhuma não conformidade registrada."
    return f"Total de NCs: {len(ncs)}"


def _list_items(items: list[Any] | None, key: str = "descricao") -> list[str]:
    if not items:
        return ["Nenhum registro."]
    out: list[str] = []
    for item in items:
        if isinstance(item, dict):
            out.append(str(item.get(key) or item))
        else:
            out.append(str(item))
    return out


def _conclusao(scores: dict[str, float]) -> str:
    geral = scores.get("conformidade_geral", 0)
    if geral >= 80:
        return f"Projeto com conformidade satisfatória ({geral}/100)."
    if geral >= 60:
        return f"Projeto com pendências moderadas ({geral}/100). Correções recomendadas."
    return f"Projeto com não conformidades relevantes ({geral}/100). Revisão obrigatória."


def _normas_text(analysis: dict[str, Any]) -> str:
    norms = analysis.get("normas_consultadas") or []
    if not norms:
        return "Normas técnicas ABNT e regulamentações aplicáveis à disciplina."
    return "; ".join(str(n.get("norma")) for n in norms[:5])


def _conclusao_parecer(analysis: dict[str, Any]) -> str:
    ncs = len(analysis.get("nao_conformidades") or [])
    if ncs == 0:
        return "Documentação analisada sem ressalvas críticas."
    return f"Foram identificadas {ncs} não conformidade(s). Parecer condicionado às correções."


def build_photographic_report_docx(
    *,
    project_name: str,
    analyses: list[dict[str, Any]],
    summary: dict[str, Any],
    obra_info: str = "",
) -> bytes:
    """Relatório fotográfico de acompanhamento de obra."""
    doc = Document()
    _add_header(doc, f"Relatório Fotográfico — {project_name}")
    doc.add_paragraph(f"Data de emissão: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    if obra_info:
        doc.add_paragraph(obra_info)

    doc.add_heading("1. Resumo executivo", level=1)
    doc.add_paragraph(
        f"Total de registros fotográficos analisados: {summary.get('analyzed', 0)} "
        f"de {summary.get('total', 0)}."
    )
    for rec in summary.get("recomendacoes") or []:
        doc.add_paragraph(f"• {rec}", style="List Bullet")

    doc.add_heading("2. Registro fotográfico", level=1)
    for idx, row in enumerate(analyses, start=1):
        if row.get("skipped") or row.get("error"):
            continue
        data = row.get("analysis") or {}
        doc.add_heading(f"Foto {idx:02d} — {row.get('filename', '—')}", level=2)
        legenda = (
            data.get("legenda_relatorio")
            or data.get("legenda_sugerida")
            or data.get("resumo_tecnico")
            or "—"
        )
        doc.add_paragraph(f"Legenda: {legenda}")
        doc.add_paragraph(f"Local: {data.get('local') or data.get('local_aproximado') or '—'}")
        doc.add_paragraph(f"Situação: {data.get('situacao') or data.get('qualidade_execucao') or '—'}")
        if data.get("observacoes_fiscal"):
            doc.add_paragraph("Observações do fiscal:")
            for obs in data["observacoes_fiscal"]:
                doc.add_paragraph(f"• {obs}", style="List Bullet")
        if data.get("acao_recomendada"):
            doc.add_paragraph(f"Ação recomendada: {data['acao_recomendada']}")

    doc.add_heading("3. Conclusão", level=1)
    ncs = summary.get("nao_conformidades") or []
    if ncs:
        doc.add_paragraph(f"Foram identificadas {len(ncs)} não conformidade(s) nas imagens analisadas.")
    else:
        doc.add_paragraph("Não foram identificadas não conformidades críticas nas imagens analisadas.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_site_inspection_laudo_docx(
    *,
    project_name: str,
    analyses: list[dict[str, Any]],
    summary: dict[str, Any],
    solicitante: str = "",
    objeto: str = "",
) -> bytes:
    """Laudo técnico consolidado a partir de vistoria fotográfica."""
    doc = Document()
    _add_header(doc, f"Laudo Técnico de Vistoria — {project_name}")
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
    if solicitante:
        doc.add_paragraph(f"Solicitante: {solicitante}")
    if objeto:
        doc.add_paragraph(f"Objeto: {objeto}")

    doc.add_heading("1. Introdução", level=1)
    doc.add_paragraph(
        "Laudo elaborado com base em análise visual automatizada (IA) das imagens "
        "vinculadas ao projeto, complementada por engenharia civil."
    )

    doc.add_heading("2. Metodologia", level=1)
    doc.add_paragraph(
        "Análise multimodal de imagens (Ollama VL) com classificação de patologias, "
        "condições observadas, não conformidades e recomendações técnicas."
    )

    doc.add_heading("3. Vistorias analisadas", level=1)
    for idx, row in enumerate(analyses, start=1):
        if row.get("skipped") or row.get("error"):
            continue
        data = row.get("analysis") or {}
        doc.add_heading(f"3.{idx} {row.get('filename', 'Imagem')}", level=2)
        doc.add_paragraph(data.get("conclusao_parcial") or data.get("resumo_tecnico") or "—")
        for pat in data.get("patologias") or []:
            if isinstance(pat, dict):
                doc.add_paragraph(
                    f"• {pat.get('tipo', 'Patologia')}: {pat.get('descricao', '')} "
                    f"({pat.get('severidade', '—')})"
                )
        if data.get("recomendacoes"):
            doc.add_paragraph("Recomendações:")
            for rec in data["recomendacoes"]:
                doc.add_paragraph(f"• {rec}", style="List Bullet")

    doc.add_heading("4. Síntese", level=1)
    doc.add_paragraph(
        f"Imagens analisadas: {summary.get('analyzed', 0)}. "
        f"NCs identificadas: {len(summary.get('nao_conformidades') or [])}."
    )
    if summary.get("normas_aplicaveis"):
        doc.add_paragraph("Normas referenciadas: " + "; ".join(summary["normas_aplicaveis"][:10]))

    doc.add_heading("5. Conclusão", level=1)
    urgent = [
        (row.get("analysis") or {}).get("urgencia")
        for row in analyses
        if row.get("analysis")
    ]
    if any(u in ("alta", "critica") for u in urgent):
        doc.add_paragraph(
            "Conclusão: foram identificadas condições de urgência alta que demandam "
            "intervenção técnica imediata."
        )
    else:
        doc.add_paragraph(
            "Conclusão: vistoria registrada. Recomenda-se acompanhamento das ações "
            "corretivas indicadas neste laudo."
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_correcoes_report_docx(
    *,
    project_name: str,
    analyses: list[dict[str, Any]],
    summary: dict[str, Any],
    prazo: str = "",
) -> bytes:
    """Relatório de correções para retorno ao projetista."""
    doc = Document()
    _add_header(doc, f"Relatório de Correções — {project_name}")
    doc.add_paragraph(f"Projeto: {project_name}")
    if prazo:
        doc.add_paragraph(f"Prazo para correção: {prazo}")

    doc.add_heading("1. Sumário", level=1)
    doc.add_paragraph(
        f"Total de arquivos analisados: {summary.get('total', 0)}. "
        f"NCs identificadas: {len(summary.get('nao_conformidades') or [])}."
    )

    doc.add_heading("2. Lista de correções", level=1)
    ncs: list[str] = list(summary.get("nao_conformidades") or [])
    if not ncs:
        for row in analyses:
            tech = row.get("technical_report") or {}
            for nc in tech.get("nao_conformidades") or []:
                if isinstance(nc, dict):
                    ncs.append(nc.get("descricao") or str(nc))
                else:
                    ncs.append(str(nc))

    if not ncs:
        doc.add_paragraph("Nenhuma correção obrigatória identificada na análise visual.")
    else:
        for idx, nc in enumerate(ncs, start=1):
            doc.add_paragraph(f"{idx}. {nc}")

    doc.add_heading("3. Recomendações", level=1)
    for rec in summary.get("recomendacoes") or []:
        doc.add_paragraph(f"• {rec}", style="List Bullet")

    doc.add_heading("4. Retorno do projetista", level=1)
    doc.add_paragraph("Espaço reservado para manifestação e comprovação das correções executadas.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_vision_technical_report_docx(
    *,
    project_name: str,
    analyses: list[dict[str, Any]],
    summary: dict[str, Any],
) -> bytes:
    """Consolida relatórios técnicos Qwen3 por arquivo analisado."""
    doc = Document()
    _add_header(doc, f"Relatório Técnico Consolidado — {project_name}")

    doc.add_heading("Resumo executivo", level=1)
    doc.add_paragraph(
        f"Analisadores: {', '.join(summary.get('analyzers_used') or []) or '—'}. "
        f"Registros: {summary.get('analyzed', 0)}/{summary.get('total', 0)}."
    )

    for idx, row in enumerate(analyses, start=1):
        if row.get("skipped") or row.get("error"):
            continue
        tech = row.get("technical_report") or {}
        if not tech:
            continue
        doc.add_heading(f"{idx}. {row.get('filename', 'Arquivo')}", level=2)
        doc.add_paragraph(tech.get("resumo_executivo") or tech.get("conclusao") or "—")
        if tech.get("parecer"):
            doc.add_paragraph(f"Parecer: {tech['parecer']}")
        for nc in tech.get("nao_conformidades") or []:
            if isinstance(nc, dict):
                doc.add_paragraph(
                    f"• NC {nc.get('codigo', '—')}: {nc.get('descricao', '')} "
                    f"(Norma: {nc.get('norma', '—')})"
                )
            else:
                doc.add_paragraph(f"• {nc}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

