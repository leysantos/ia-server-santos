"""Exportação DOCX do laudo — formatação técnica institucional.

Cabeçalho compacto: logo (esq.) + empresa; dados (dir.); linhas azul/cinza.
Rodapé: linhas azul/cinza + endereço + Página X de Y.
Brasão: marca d'água central grande (atrás do texto).
Corpo: Times New Roman justificado, capítulos numerados, cards/gráficos, fotos sem páginas vazias.
"""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from core.inspection_report.analytics import fit_image_display_inches, prepare_watermark_png
from core.inspection_report.format_utils import (
    COLOR_BLUE,
    COLOR_GRAY,
    build_body_sections,
    build_cover_layout,
    build_photographic_presentation,
    build_sumario_entries,
    format_generated_at,
    header_meta_lines,
    inject_coordinates_into_object_tables,
    normalize_parties,
    party_display_lines,
    photo_source_line,
)
from core.system.company_profile import (
    get_company_profile,
    load_company_brasao,
    load_company_logo,
)

FIRST_LINE_INDENT = Cm(1.5)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
GRAY = RGBColor(0x47, 0x55, 0x69)
BLACK = RGBColor(0x0F, 0x17, 0x2A)


def _set_run_font(run, *, size=11, bold=False, color: RGBColor | None = None):
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            element.set(qn(f"w:{k}"), str(v))
        tcBorders.append(element)


def _set_paragraph_borders(paragraph, *, top=None, bottom=None):
    pPr = paragraph._p.get_or_add_pPr()
    for child in list(pPr):
        if child.tag == qn("w:pBdr"):
            pPr.remove(child)
    pBdr = OxmlElement("w:pBdr")
    if top:
        el = OxmlElement("w:top")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(top[1]))
        el.set(qn("w:space"), "1")
        el.set(qn("w:color"), top[0])
        pBdr.append(el)
    if bottom:
        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(bottom[1]))
        el.set(qn("w:space"), "1")
        el.set(qn("w:color"), bottom[0])
        pBdr.append(el)
    pPr.append(pBdr)


def _add_dual_separator_paragraphs(container, *, as_header: bool):
    if as_header:
        p1 = container.add_paragraph()
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(0)
        _set_paragraph_borders(p1, bottom=(COLOR_BLUE, "18"))
        p2 = container.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(1)
        _set_paragraph_borders(p2, bottom=(COLOR_GRAY, "10"))
    else:
        p1 = container.add_paragraph()
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(0)
        _set_paragraph_borders(p1, top=(COLOR_BLUE, "18"))
        p2 = container.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(1)
        _set_paragraph_borders(p2, top=(COLOR_GRAY, "10"))


def _clear_paragraph(p):
    p.clear()


def _add_field(paragraph, instr_text: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    _set_run_font(run, size=8)


def _set_table_fixed(table, widths_cm: list[float]):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def _nil_borders(cell):
    for edge in ("top", "left", "bottom", "right"):
        _set_cell_border(
            cell,
            **{edge: {"val": "nil", "sz": "0", "color": "FFFFFF", "space": "0"}},
        )


def _add_watermark(section, brasao_bytes: bytes):
    """Brasão como marca d'água na largura do corpo da página (atrás do texto)."""
    # Corpo A4: 21 − 2,5 − 2,0 = 16,5 cm
    body_w_cm = 16.5
    wm = prepare_watermark_png(
        brasao_bytes,
        size_px=1800,
        opacity=0.12,
        max_width_px=1800,
    )
    if not wm:
        return
    header = section.header
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    try:
        inline_shape = run.add_picture(io.BytesIO(wm), width=Cm(body_w_cm))
    except Exception:
        return

    # Converte wp:inline → wp:anchor (behindDoc=1), centralizado na página
    inline = inline_shape._inline
    drawing = inline.getparent()
    if drawing is None:
        return

    extent = inline.find(qn("wp:extent"))
    docPr = inline.find(qn("wp:docPr"))
    graphic = inline.find(qn("a:graphic"))
    if extent is None or graphic is None:
        return

    cx = extent.get("cx")
    cy = extent.get("cy")
    doc_id = docPr.get("id") if docPr is not None else "1"
    doc_name = docPr.get("name") if docPr is not None else "Watermark"

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "0")
    anchor.set("behindDoc", "1")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    # Centro horizontal da página
    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    align_h = OxmlElement("wp:align")
    align_h.text = "center"
    pos_h.append(align_h)
    anchor.append(pos_h)

    # Centro vertical da página
    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    align_v = OxmlElement("wp:align")
    align_v.text = "center"
    pos_v.append(align_v)
    anchor.append(pos_v)

    ext2 = OxmlElement("wp:extent")
    ext2.set("cx", cx)
    ext2.set("cy", cy)
    anchor.append(ext2)

    wrap = OxmlElement("wp:wrapNone")
    anchor.append(wrap)

    doc_pr = OxmlElement("wp:docPr")
    doc_pr.set("id", str(doc_id))
    doc_pr.set("name", str(doc_name))
    anchor.append(doc_pr)

    cNv = OxmlElement("wp:cNvGraphicFramePr")
    anchor.append(cNv)

    anchor.append(graphic)
    drawing.replace(inline, anchor)


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    first_indent: bool = True,
    bold: bool = False,
    size: int = 11,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after: int = 8,
    color: RGBColor | None = None,
):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_indent and text and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = FIRST_LINE_INDENT
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)
    return p


def _setup_header_footer(doc: Document, *, content: dict[str, Any], generated_at: str):
    company = get_company_profile()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # Cabeçalho compacto
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.35)
    section.footer_distance = Cm(0.45)

    org = company.display_name() or company.razao_social or "Empresa responsável"
    numero = str(content.get("numero_laudo") or content.get("titulo") or "Laudo técnico")

    header = section.header
    header.is_linked_to_previous = False
    while len(header.paragraphs) > 1:
        header.paragraphs[-1]._element.getparent().remove(header.paragraphs[-1]._element)
    _clear_paragraph(header.paragraphs[0])
    header.paragraphs[0].paragraph_format.space_after = Pt(0)

    ht = header.add_table(rows=1, cols=2, width=Cm(16.5))
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed(ht, [8.0, 8.5])
    left, right = ht.rows[0].cells
    _nil_borders(left)
    _nil_borders(right)

    # Logo + nome na mesma célula, compacto
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_after = Pt(0)
    logo = load_company_logo()
    if logo:
        try:
            lp.add_run().add_picture(io.BytesIO(logo), width=Inches(0.72))
            lp.add_run("  ")
        except Exception:
            pass
    run = lp.add_run(org.upper())
    _set_run_font(run, size=8, bold=True, color=BLUE)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    meta = header_meta_lines(content, generated_at=generated_at)
    for i, line in enumerate(meta):
        p = rp if i == 0 else right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        _set_run_font(run, size=7, bold=(i == 0), color=GRAY)

    _add_dual_separator_paragraphs(header, as_header=True)

    # Marca d'água (brasão)
    brasao = load_company_brasao()
    if brasao:
        _add_watermark(section, brasao)

    footer = section.footer
    footer.is_linked_to_previous = False
    while len(footer.paragraphs) > 1:
        footer.paragraphs[-1]._element.getparent().remove(footer.paragraphs[-1]._element)
    _clear_paragraph(footer.paragraphs[0])
    _add_dual_separator_paragraphs(footer, as_header=False)

    addr = " | ".join(
        x for x in (company.endereco_linha() or None, company.site or company.email or None) if x
    ) or org
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(1)
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run(addr)
    _set_run_font(run, size=7, color=GRAY)

    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp2.paragraph_format.space_before = Pt(0)
    run = fp2.add_run("Página ")
    _set_run_font(run, size=8)
    _add_field(fp2, "PAGE")
    run = fp2.add_run(" de ")
    _set_run_font(run, size=8)
    _add_field(fp2, "NUMPAGES")
    run = fp2.add_run(f"  |  {numero}"[:90])
    _set_run_font(run, size=7, color=GRAY)


def _add_table(doc: Document, table_data: dict[str, Any]):
    headers = table_data.get("headers") or []
    rows = table_data.get("rows") or []
    if not headers:
        return
    if table_data.get("caption"):
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(4)
        run = cap.add_run(str(table_data["caption"]))
        _set_run_font(run, size=10, bold=True, color=BLACK)

    ncols = len(headers)
    if ncols == 8:
        widths = [1.0, 1.3, 3.0, 2.4, 1.8, 1.2, 1.4, 2.4]
    elif ncols == 7:
        widths = [1.3, 3.0, 2.4, 1.7, 1.2, 4.0, 2.4]
    elif ncols == 4:
        widths = [2.5, 2.5, 2.5, 9.0]
    elif ncols == 2:
        widths = [5.5, 11.0]
    else:
        widths = [16.5 / ncols] * ncols

    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = "Table Grid"
    _set_table_fixed(tbl, widths)

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        _set_run_font(run, size=8, bold=True)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E2E8F0")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    for r_idx, row in enumerate(rows):
        for c_idx in range(ncols):
            val = row[c_idx] if c_idx < len(row) else ""
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            _set_run_font(run, size=8)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_cards(doc: Document, cards: list[dict[str, Any]]):
    if not cards:
        return
    # 3 colunas por linha
    chunk = 3
    for i in range(0, len(cards), chunk):
        group = cards[i : i + chunk]
        cols = len(group)
        tbl = doc.add_table(rows=1, cols=cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        width = 16.5 / cols
        _set_table_fixed(tbl, [width] * cols)
        for ci, card in enumerate(group):
            cell = tbl.rows[0].cells[ci]
            _nil_borders(cell)
            # borda suave
            for edge in ("top", "left", "bottom", "right"):
                _set_cell_border(
                    cell,
                    **{edge: {"val": "single", "sz": "8", "color": "CBD5E1", "space": "0"}},
                )
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F8FAFC")
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)

            p0 = cell.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p0.add_run(str(card.get("label") or ""))
            _set_run_font(run, size=8, bold=True, color=GRAY)

            p1 = cell.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p1.add_run(str(card.get("value") or ""))
            _set_run_font(run, size=12, bold=True, color=BLUE)

            if card.get("hint"):
                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p2.add_run(str(card["hint"]))
                _set_run_font(run, size=7, color=GRAY)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _add_heading(doc: Document, text: str, level: int = 1, *, align=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    for run in h.runs:
        size = 14 if level == 0 else (12 if level == 1 else 11)
        _set_run_font(run, size=size, bold=True, color=BLUE if level == 1 else BLACK)
    return h


def _add_signature_cell(cell, party: dict[str, Any]) -> None:
    """Preenche célula de assinatura com tudo centralizado."""
    _nil_borders(cell)
    # limpa parágrafo padrão
    p0 = cell.paragraphs[0]
    p0.text = ""
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(28)
    p0.paragraph_format.space_after = Pt(4)
    run = p0.add_run("_______________________________")
    _set_run_font(run, size=10, color=BLACK)

    lines = party_display_lines(party)
    for idx, line in enumerate(lines):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.right_indent = Pt(0)
        r = p.add_run(line)
        _set_run_font(r, size=10, bold=(idx == 0), color=BLACK)


def _add_signatures_block(doc: Document, content: dict[str, Any]) -> None:
    """Assinaturas dos responsáveis técnicos antes do relatório fotográfico."""
    parties = normalize_parties(content.get("responsaveis_tecnicos"))
    if not parties:
        return
    doc.add_page_break()
    _add_heading(doc, "Responsáveis técnicos", level=1)
    _add_paragraph(
        doc,
        "Declaramos, para os devidos fins, a responsabilidade técnica pelo presente laudo, "
        "nos termos das normas aplicáveis e do registro profissional abaixo indicado.",
        first_indent=True,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=16,
    )
    # Um responsável: bloco centralizado sem tabela (evita desalinhamento)
    if len(parties) == 1:
        _add_paragraph(
            doc,
            "_______________________________",
            first_indent=False,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=4,
            size=10,
        )
        for idx, line in enumerate(party_display_lines(parties[0])):
            _add_paragraph(
                doc,
                line,
                first_indent=False,
                bold=(idx == 0),
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=1,
            )
        return

    for i in range(0, len(parties), 2):
        chunk = parties[i : i + 2]
        widths = [8.0, 8.0] if len(chunk) == 2 else [16.5]
        table = doc.add_table(rows=1, cols=len(chunk))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_fixed(table, widths[: len(chunk)])
        for cell, party in zip(table.rows[0].cells, chunk):
            _add_signature_cell(cell, party)
        doc.add_paragraph().paragraph_format.space_after = Pt(10)


def _add_kv_table(doc: Document, rows: list[list[str]], *, label_w: float = 4.8, value_w: float = 11.7):
    """Tabela rótulo|valor da capa — rótulo em negrito com fundo suave."""
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    _set_table_fixed(tbl, [label_w, value_w])
    for r_idx, row in enumerate(rows):
        label = str(row[0] if row else "")
        value = str(row[1] if len(row) > 1 else "")
        c0 = tbl.rows[r_idx].cells[0]
        c1 = tbl.rows[r_idx].cells[1]
        c0.text = ""
        c1.text = ""
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        _set_run_font(r0, size=9, bold=True, color=BLACK)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8EEF9")
        shading.set(qn("w:val"), "clear")
        c0._tc.get_or_add_tcPr().append(shading)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(value)
        _set_run_font(r1, size=10, bold=False, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def _add_cover_section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text.upper())
    _set_run_font(run, size=10, bold=True, color=RGBColor(0x1D, 0x4E, 0xD8))


def _add_cover_page(doc: Document, content: dict[str, Any], *, generated_at: str) -> None:
    """1ª folha institucional — título + blocos em tabela bem distribuídos."""
    cover = build_cover_layout(content, generated_at=generated_at)

    _add_paragraph(
        doc,
        str(cover["titulo"]).upper(),
        first_indent=False,
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        color=BLACK,
    )
    if cover.get("subtitulo"):
        _add_paragraph(
            doc,
            str(cover["subtitulo"]),
            first_indent=False,
            bold=True,
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=6,
        )

    # Linha divisória sob o título
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(10)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLOR_BLUE)
    pBdr.append(bottom)
    pPr.append(pBdr)

    for block in cover.get("blocks") or []:
        heading = str(block.get("heading") or "").strip()
        rows = block.get("rows") or []
        if not rows:
            continue
        if heading:
            _add_cover_section_title(doc, heading)
        _add_kv_table(doc, rows)

    _add_paragraph(
        doc,
        f"Documento gerado em: {generated_at}",
        first_indent=False,
        size=9,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=8,
        color=GRAY,
    )

    note = cover.get("compliance_note") or ""
    if note:
        _add_cover_section_title(doc, "Conformidade normativa")
        _add_paragraph(
            doc,
            str(note),
            first_indent=False,
            size=9,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=6,
        )


def _add_sumario(doc: Document, content: dict[str, Any]) -> None:
    """Insere página de Sumário após a capa, com a lista de seções do laudo."""
    entries = build_sumario_entries(content)
    if not entries:
        return
    doc.add_page_break()
    _add_heading(doc, "Sumário", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(
        doc,
        "O presente laudo está estruturado conforme as seções abaixo.",
        first_indent=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10,
        space_after=10,
    )
    for idx, entry in enumerate(entries, start=1):
        label = (entry.get("label") or "").strip()
        if not label:
            continue
        # Se o título já tem numeração (ex.: "1. Solicitação"), não prefixa índice do sumário
        line = label if label[:1].isdigit() else f"{idx}. {label}"
        _add_paragraph(
            doc,
            line,
            first_indent=False,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            size=11,
            space_after=4,
        )


def build_inspection_laudo_docx(
    *,
    content: dict[str, Any],
    image_assets: list[dict[str, Any]],
    georef_asset: dict[str, Any] | None = None,
) -> bytes:
    generated_at = format_generated_at(datetime.now())
    doc = Document()

    export_content = dict(content or {})
    if georef_asset and georef_asset.get("latitude") is not None:
        export_content = inject_coordinates_into_object_tables(
            export_content,
            latitude=georef_asset.get("latitude"),
            longitude=georef_asset.get("longitude"),
            label=georef_asset.get("label"),
        )

    _setup_header_footer(doc, content=export_content, generated_at=generated_at)

    _add_cover_page(doc, export_content, generated_at=generated_at)

    _add_sumario(doc, export_content)

    sections = build_body_sections(export_content)
    for section in sections:
        _add_heading(doc, section["title"], level=1)
        if section.get("cards"):
            _add_cards(doc, section["cards"])
        for para in section.get("paragraphs") or []:
            _add_paragraph(doc, str(para), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for table in section.get("tables") or []:
            _add_table(doc, table)

        # Imagem georreferenciada logo abaixo da tabela de dados técnicos do objeto
        cid = str(section.get("chapter_id") or "")
        title_l = str(section.get("title") or "").lower()
        is_ficha = cid == "ficha_tecnica" or "ficha técnica" in title_l or "ficha tecnica" in title_l
        if is_ficha and georef_asset and georef_asset.get("path") and Path(georef_asset["path"]).exists():
            try:
                dw, dh = fit_image_display_inches(str(georef_asset["path"]), max_w=5.9, max_h=4.2)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                p.add_run().add_picture(
                    str(georef_asset["path"]), width=Inches(dw), height=Inches(dh)
                )
                cap = georef_asset.get("caption") or "Imagem georreferenciada do objeto"
                if georef_asset.get("label"):
                    cap = f"{cap} — {georef_asset['label']}"
                _add_paragraph(
                    doc,
                    cap,
                    first_indent=False,
                    size=9,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=8,
                )
            except Exception:
                _add_paragraph(
                    doc,
                    "[Falha ao inserir imagem georreferenciada]",
                    first_indent=False,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                )

        for chart in section.get("charts") or []:
            labels = chart.get("labels") or []
            values = chart.get("values") or []
            _add_table(
                doc,
                {
                    "caption": chart.get("caption") or "Dados do gráfico",
                    "headers": ["Categoria", "Valor"],
                    "rows": [[str(l), str(v)] for l, v in zip(labels, values)],
                },
            )
        for ch_img in section.get("chart_images") or []:
            if ch_img.get("caption"):
                _add_paragraph(
                    doc,
                    str(ch_img["caption"]),
                    first_indent=False,
                    bold=True,
                    size=10,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_after=4,
                )
            png = ch_img.get("png")
            if png:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(io.BytesIO(png), width=Inches(5.9))
                except Exception:
                    pass

    # Assinaturas dos RT — imediatamente antes do relatório fotográfico
    _add_signatures_block(doc, export_content)

    # Relatório fotográfico
    photo_num = (sections[-1]["number"] + 1) if sections else 1
    doc.add_page_break()
    _add_heading(doc, f"{photo_num}. Relatório fotográfico", level=1)
    _add_paragraph(
        doc,
        build_photographic_presentation(export_content),
        first_indent=True,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    photo_entries = export_content.get("photographic_report") or []
    path_by_file = {a["filename"].lower(): a for a in image_assets}
    path_by_num = {int(a["photo_number"]): a for a in image_assets if a.get("photo_number")}
    ordered = sorted(photo_entries, key=lambda p: int(p.get("photo_number") or 0))
    if not ordered:
        ordered = [
            {
                "photo_number": a.get("photo_number"),
                "filename": a["filename"],
                "title": f"Foto {a.get('photo_number') or '':02d}",
                "description": a.get("caption") or "Registro fotográfico.",
                "legend": "",
            }
            for a in image_assets
        ]

    for i, entry in enumerate(ordered):
        # Uma quebra por foto; conteúdo compacto para não gerar página órfã
        doc.add_page_break()
        num = int(entry.get("photo_number") or 0)
        heading = f"Foto {num:02d} – {entry.get('title') or entry.get('filename') or ''}"
        _add_heading(doc, heading, level=2, align=WD_ALIGN_PARAGRAPH.CENTER)

        asset = path_by_num.get(num) or path_by_file.get(str(entry.get("filename") or "").lower())
        if asset and Path(asset["path"]).exists():
            try:
                dw, dh = fit_image_display_inches(str(asset["path"]), max_w=5.9, max_h=5.0)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)
                p.add_run().add_picture(str(asset["path"]), width=Inches(dw), height=Inches(dh))
            except Exception:
                _add_paragraph(
                    doc,
                    f"[Falha ao inserir imagem: {asset['filename']}]",
                    first_indent=False,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                )

        _add_paragraph(
            doc,
            f"Descrição: {entry.get('description') or '—'}",
            first_indent=True,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=4,
        )
        if entry.get("legend"):
            _add_paragraph(
                doc,
                f"Legenda: {entry.get('legend')}",
                first_indent=False,
                size=9,
                bold=True,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_after=2,
            )
        _add_paragraph(
            doc,
            photo_source_line(export_content),
            first_indent=False,
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=2,
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
