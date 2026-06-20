"""
Extração de texto de arquivos de projeto para RAG (multi-formato).
"""

from __future__ import annotations

import csv
import json
import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)

# Sincronizado com upload/indexação do workspace
PROJECT_INDEXABLE_SUFFIXES: frozenset[str] = frozenset(
    {
        ".pdf",
        ".txt",
        ".md",
        ".json",
        ".csv",
        ".xlsx",
        ".xls",
        ".docx",
        ".dxf",
        ".dwg",
        ".ifc",
        ".rtf",
        ".png",
        ".jpg",
        ".jpeg",
        ".webp",
        ".heic",
        ".heif",
        ".zip",
    }
)

PROJECT_UPLOAD_ACCEPT = ",".join(
    sorted(
        {
            ".pdf",
            ".txt",
            ".md",
            ".csv",
            ".xlsx",
            ".xls",
            ".docx",
            ".dxf",
            ".dwg",
            ".ifc",
            ".json",
            ".rtf",
            ".png",
            ".jpg",
            ".jpeg",
            ".webp",
            ".heic",
            ".heif",
            ".zip",
            "image/png",
            "image/jpeg",
            "image/webp",
            "image/heic",
            "application/zip",
            "application/pdf",
            "text/plain",
            "text/csv",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    )
)


@dataclass
class ExtractedSegment:
    """Trecho indexável — section label aparece no contexto RAG."""

    text: str
    section: str = ""
    section_num: int = 0


def is_indexable_project_file(path: Path | str) -> bool:
    return Path(path).suffix.lower() in PROJECT_INDEXABLE_SUFFIXES


def extract_project_file_segments(path: Path | str) -> tuple[list[ExtractedSegment], str]:
    """
    Extrai segmentos de texto do arquivo.
    Returns: (segments, format_key)
    Raises: ValueError for unsupported/empty; ImportError when optional dep missing.
    """
    file_path = Path(path).resolve()
    suffix = file_path.suffix.lower()
    extractors: dict[str, Callable[[Path], tuple[list[ExtractedSegment], str]]] = {
        ".pdf": _extract_pdf,
        ".txt": _extract_plain_text,
        ".md": _extract_plain_text,
        ".json": _extract_json,
        ".csv": _extract_csv,
        ".xlsx": _extract_excel,
        ".xls": _extract_xls,
        ".docx": _extract_docx,
        ".dxf": _extract_dxf,
        ".dwg": _extract_dwg,
        ".ifc": _extract_ifc,
        ".rtf": _extract_rtf,
        ".png": _extract_image,
        ".jpg": _extract_image,
        ".jpeg": _extract_image,
        ".zip": _extract_zip,
    }

    if suffix not in extractors:
        raise ValueError(f"Formato não suportado para indexação: {suffix}")

    segments, fmt = extractors[suffix](file_path)
    segments = [s for s in segments if s.text.strip()]
    if not segments:
        raise ValueError(f"Nenhum texto extraído de {file_path.name}")
    return segments, fmt


def _extract_pdf(path: Path) -> tuple[list[ExtractedSegment], str]:
    from memory.pdf_indexer import PDFIndexer

    pages = PDFIndexer.extract_text(path)
    return (
        [
            ExtractedSegment(text=text, section=f"p.{num}", section_num=num)
            for num, text in pages
        ],
        "pdf",
    )


def _extract_plain_text(path: Path) -> tuple[list[ExtractedSegment], str]:
    raw = _read_text_file(path)
    return [ExtractedSegment(text=raw, section="documento", section_num=1)], path.suffix.lstrip(".")


def _extract_json(path: Path) -> tuple[list[ExtractedSegment], str]:
    raw = _read_text_file(path)
    try:
        parsed = json.loads(raw)
        pretty = json.dumps(parsed, ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        pretty = raw
    return [ExtractedSegment(text=pretty, section="json", section_num=1)], "json"


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_csv(path: Path) -> tuple[list[ExtractedSegment], str]:
    segments: list[ExtractedSegment] = []
    with open(path, newline="", encoding="utf-8-sig") as handle:
        reader = csv.DictReader(handle)
        for row_num, row in enumerate(reader, start=2):
            text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
            if text.strip():
                segments.append(
                    ExtractedSegment(text=text, section=f"linha {row_num}", section_num=row_num)
                )
    return segments, "csv"


def _extract_excel(path: Path) -> tuple[list[ExtractedSegment], str]:
    try:
        import openpyxl
    except ImportError as exc:
        raise ImportError("openpyxl necessário: pip install openpyxl") from exc

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    segments: list[ExtractedSegment] = []
    for sheet_idx, sheet_name in enumerate(wb.sheetnames, start=1):
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h or f"col{i}") for i, h in enumerate(rows[0])]
        for row_num, row in enumerate(rows[1:], start=2):
            cells = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
            text = " | ".join(f"{k}: {v}" for k, v in cells.items() if v is not None)
            if text.strip():
                segments.append(
                    ExtractedSegment(
                        text=text,
                        section=f"aba:{sheet_name}",
                        section_num=row_num + sheet_idx * 10_000,
                    )
                )
    wb.close()
    return segments, "xlsx"


def _extract_xls(path: Path) -> tuple[list[ExtractedSegment], str]:
    try:
        import xlrd
    except ImportError as exc:
        raise ImportError("xlrd necessário para .xls: pip install xlrd") from exc

    book = xlrd.open_workbook(str(path))
    segments: list[ExtractedSegment] = []
    for sheet_idx, sheet in enumerate(book.sheets()):
        if sheet.nrows == 0:
            continue
        headers = [str(sheet.cell_value(0, c) or f"col{c}") for c in range(sheet.ncols)]
        for row_num in range(1, sheet.nrows):
            cells = {
                headers[c]: sheet.cell_value(row_num, c)
                for c in range(min(sheet.ncols, len(headers)))
            }
            text = " | ".join(f"{k}: {v}" for k, v in cells.items() if v not in ("", None))
            if str(text).strip():
                segments.append(
                    ExtractedSegment(
                        text=text,
                        section=f"aba:{sheet.name}",
                        section_num=row_num + sheet_idx * 10_000,
                    )
                )
    return segments, "xls"


def _extract_docx(path: Path) -> tuple[list[ExtractedSegment], str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx necessário: pip install python-docx") from exc

    doc = Document(str(path))
    segments: list[ExtractedSegment] = []
    para_num = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        para_num += 1
        segments.append(
            ExtractedSegment(text=text, section=f"§{para_num}", section_num=para_num)
        )

    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not cells:
                continue
            text = " | ".join(cells)
            segments.append(
                ExtractedSegment(
                    text=text,
                    section=f"tabela{table_idx}:linha{row_idx}",
                    section_num=table_idx * 1000 + row_idx,
                )
            )
    return segments, "docx"


def _extract_dxf(path: Path) -> tuple[list[ExtractedSegment], str]:
    try:
        import ezdxf
    except ImportError as exc:
        raise ImportError("ezdxf necessário para DXF: pip install ezdxf") from exc

    doc = ezdxf.readfile(str(path))
    msp = doc.modelspace()
    segments: list[ExtractedSegment] = []

    layers = sorted({layer.dxf.name for layer in doc.layers})
    if layers:
        segments.append(
            ExtractedSegment(
                text="Camadas DXF: " + ", ".join(layers[:80]),
                section="camadas",
                section_num=1,
            )
        )

    texts: list[str] = []
    for entity in msp:
        if entity.dxftype() in ("TEXT", "MTEXT"):
            try:
                content = entity.dxf.text if entity.dxftype() == "TEXT" else entity.text
                if content and str(content).strip():
                    texts.append(str(content).strip())
            except Exception:
                continue

    for idx, text in enumerate(texts[:200], start=1):
        segments.append(ExtractedSegment(text=text, section=f"texto CAD", section_num=idx))

    if not segments:
        segments.append(
            ExtractedSegment(
                text=f"Arquivo DXF {path.name} — sem entidades TEXT/MTEXT legíveis.",
                section="meta",
                section_num=0,
            )
        )
    return segments, "dxf"


def _extract_dwg(path: Path) -> tuple[list[ExtractedSegment], str]:
    """
    DWG é binário proprietário — extrai strings ASCII legíveis + metadados.
    Para indexação completa, exporte para PDF ou DXF.
    """
    data = path.read_bytes()
    strings = _extract_printable_strings(data, min_len=5)
    unique = list(dict.fromkeys(strings))[:300]

    segments: list[ExtractedSegment] = [
        ExtractedSegment(
            text=(
                f"Arquivo DWG: {path.name} ({len(data) // 1024} KB). "
                "Texto parcial extraído de strings internas; "
                "para melhor RAG exporte para PDF ou DXF."
            ),
            section="meta",
            section_num=0,
        )
    ]
    for idx, text in enumerate(unique, start=1):
        if len(text) > 4000:
            text = text[:4000] + "…"
        segments.append(ExtractedSegment(text=text, section="string CAD", section_num=idx))
    return segments, "dwg"


def _extract_printable_strings(data: bytes, min_len: int = 4) -> list[str]:
    pattern = re.compile(rb"[\x20-\x7e\xC0-\xFF]{%d,}" % min_len)
    found: list[str] = []
    for match in pattern.finditer(data):
        try:
            s = match.group().decode("utf-8", errors="ignore").strip()
        except Exception:
            continue
        if len(s) >= min_len and not s.isdigit():
            found.append(s)
    return found


def _extract_ifc(path: Path) -> tuple[list[ExtractedSegment], str]:
    try:
        import ifcopenshell
        import ifcopenshell.util.element as element_util
    except ImportError as exc:
        raise ImportError("ifcopenshell necessário: pip install ifcopenshell") from exc

    model = ifcopenshell.open(str(path))
    segments: list[ExtractedSegment] = []

    project = model.by_type("IfcProject")
    if project:
        name = getattr(project[0], "Name", None) or path.stem
        desc = getattr(project[0], "Description", None) or ""
        segments.append(
            ExtractedSegment(
                text=f"Projeto IFC: {name}. {desc}".strip(),
                section="IfcProject",
                section_num=1,
            )
        )

    buildings = model.by_type("IfcBuilding")
    if buildings:
        names = [getattr(b, "Name", None) or "Edificação" for b in buildings[:20]]
        segments.append(
            ExtractedSegment(
                text="Edificações: " + ", ".join(names),
                section="IfcBuilding",
                section_num=2,
            )
        )

    storeys = model.by_type("IfcBuildingStorey")
    if storeys:
        storey_lines = []
        for s in storeys[:40]:
            label = getattr(s, "Name", None) or "Pavimento"
            elevation = getattr(s, "Elevation", None)
            if elevation is not None:
                storey_lines.append(f"{label} (cota {elevation})")
            else:
                storey_lines.append(label)
        segments.append(
            ExtractedSegment(
                text="Pavimentos: " + "; ".join(storey_lines),
                section="IfcBuildingStorey",
                section_num=3,
            )
        )

    type_counts: dict[str, int] = {}
    for element in model.by_type("IfcProduct"):
        type_counts[element.is_a()] = type_counts.get(element.is_a(), 0) + 1
    if type_counts:
        top = sorted(type_counts.items(), key=lambda x: -x[1])[:25]
        summary = ", ".join(f"{t}: {n}" for t, n in top)
        segments.append(
            ExtractedSegment(
                text=f"Elementos IFC (contagem): {summary}",
                section="IfcProduct",
                section_num=4,
            )
        )

    # Amostra de propriedades de poucos elementos estruturais
    samples = 0
    for element in model.by_type("IfcBuildingElement"):
        if samples >= 15:
            break
        try:
            psets = element_util.get_psets(element)
        except Exception:
            continue
        if not psets:
            continue
        flat = []
        for pset_name, props in psets.items():
            for key, val in list(props.items())[:8]:
                if key != "id" and val is not None:
                    flat.append(f"{pset_name}.{key}={val}")
        if flat:
            label = getattr(element, "Name", None) or element.is_a()
            segments.append(
                ExtractedSegment(
                    text=f"{label}: " + " | ".join(flat[:12]),
                    section="IfcPropertySet",
                    section_num=10 + samples,
                )
            )
            samples += 1

    if not segments:
        segments.append(
            ExtractedSegment(
                text=f"Modelo IFC {path.name} carregado sem metadados textuais extraíveis.",
                section="meta",
                section_num=0,
            )
        )
    return segments, "ifc"


def _extract_rtf(path: Path) -> tuple[list[ExtractedSegment], str]:
    raw = _read_text_file(path)
    # Remoção simples de tags RTF — suficiente para RAG básico
    text = re.sub(r"\\[a-z]+\d* ?", " ", raw)
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        raise ValueError("RTF sem texto legível")
    return [ExtractedSegment(text=text, section="rtf", section_num=1)], "rtf"


def _extract_image(path: Path) -> tuple[list[ExtractedSegment], str]:
    """Metadados básicos de imagem; OCR profundo via Project Review Engine."""
    try:
        from core.project_review.extraction.ocr_pipeline import extract_structured

        data = extract_structured(path)
        text = data.get("texto") or f"Imagem {path.name}"
        if data.get("carimbos"):
            text += "\nCarimbos: " + ", ".join(data["carimbos"][:5])
    except Exception:
        text = f"Imagem de projeto: {path.name}"
    fmt = path.suffix.lstrip(".").lower()
    return [ExtractedSegment(text=text, section="imagem", section_num=1)], fmt


def _extract_zip(path: Path) -> tuple[list[ExtractedSegment], str]:
    import zipfile

    names: list[str] = []
    combined: list[str] = []
    with zipfile.ZipFile(path, "r") as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            inner = Path(info.filename).name
            if inner:
                names.append(inner)
            if Path(inner).suffix.lower() in {".txt", ".md", ".csv"}:
                try:
                    combined.append(zf.read(info).decode("utf-8", errors="ignore")[:8000])
                except Exception:
                    pass
    text = f"Arquivo ZIP: {path.name}\nConteúdo: {', '.join(names[:40])}"
    if combined:
        text += "\n\n" + "\n---\n".join(combined[:5])
    return [ExtractedSegment(text=text, section="zip", section_num=1)], "zip"
