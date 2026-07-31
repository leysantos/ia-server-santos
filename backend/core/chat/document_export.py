"""Exporta resposta do chat como documento técnico no layout institucional dos Laudos."""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any, Literal

DocKind = Literal[
    "memoria",
    "trd",
    "memorial",
    "parecer",
    "especificacao",
    "checklist",
    "nota_orcamento",
    "resposta",
]
DocFormat = Literal["pdf", "docx"]

_TITLE = {
    "memoria": "Memória de Cálculo",
    "trd": "Termo de Referência Descritivo (TRD)",
    "memorial": "Memorial Descritivo",
    "parecer": "Parecer Técnico",
    "especificacao": "Especificação Técnica",
    "checklist": "Checklist de Verificação Técnica",
    "nota_orcamento": "Nota Técnica de Orçamento",
    "resposta": "Resposta Técnica — Chat IA",
}

_DOC_CODE = {
    "memoria": "MC",
    "trd": "TRD",
    "memorial": "MD",
    "parecer": "PT",
    "especificacao": "ET",
    "checklist": "CHK",
    "nota_orcamento": "ORC",
    "resposta": "RT",
}

_SECTIONED = frozenset({"trd", "memorial", "parecer", "especificacao", "checklist"})


def _clean_md(text: str) -> str:
    t = (text or "").replace("\r\n", "\n").strip()
    t = re.sub(r"\$([^$]+)\$", r"\1", t)
    t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", t)
    t = re.sub(r"`([^`]+)`", r"\1", t)
    return t


def _split_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    buf: list[str] = []
    for raw_line in _clean_md(text).split("\n"):
        line = raw_line.rstrip()
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            if buf:
                blocks.append(("para", "\n".join(buf).strip()))
                buf = []
            blocks.append(("heading", heading.group(2).strip()))
            continue
        if line.startswith("|") and "|" in line[1:]:
            if buf:
                blocks.append(("para", "\n".join(buf).strip()))
                buf = []
            blocks.append(("table_row", line))
            continue
        if re.match(r"^-{3,}$", line.strip()):
            if buf:
                blocks.append(("para", "\n".join(buf).strip()))
                buf = []
            continue
        buf.append(line)
    if buf:
        blocks.append(("para", "\n".join(buf).strip()))
    return [(k, v) for k, v in blocks if v]


def _clip(text: str, n: int = 6000) -> str:
    t = text.strip()
    return t[:n] + ("…" if len(t) > n else "")


def _disclaimer() -> str:
    return (
        "Documento gerado com apoio de IA. A validação final, ART e "
        "responsabilidade técnica cabem ao profissional legalmente habilitado."
    )


def _trd_sections(body: str) -> list[tuple[str, str]]:
    cleaned = _clean_md(body)
    return [
        (
            "1. Objeto",
            "Elaboração de documentação técnica de engenharia a partir da solução "
            "recomendada pelo Chat IA Server Santos.",
        ),
        ("2. Escopo técnico", _clip(cleaned)),
        (
            "3. Entregáveis",
            "• Memória de cálculo com premissas, esforços e dimensionamento\n"
            "• Detalhamento de armaduras / quantitativo de aço (quando aplicável)\n"
            "• Croqui esquemático (quando gerado)\n"
            "• Referências normativas citadas na solução",
        ),
        (
            "4. Normas de referência",
            "Normas ABNT e demais referências citadas na solução técnica anexa, "
            "além de legislação aplicável ao empreendimento.",
        ),
        ("5. Premissas e responsabilidades", _disclaimer()),
    ]


def _memorial_sections(body: str) -> list[tuple[str, str]]:
    cleaned = _clean_md(body)
    return [
        (
            "1. Objeto",
            "Memorial descritivo elaborado a partir da resposta técnica do Chat IA.",
        ),
        ("2. Descrição técnica", _clip(cleaned)),
        (
            "3. Materiais e sistemas",
            "Conforme descritos na seção anterior; complementar com catálogos "
            "e fichas técnicas do empreendimento quando disponíveis.",
        ),
        ("4. Responsabilidades", _disclaimer()),
    ]


def _parecer_sections(body: str) -> list[tuple[str, str]]:
    cleaned = _clean_md(body)
    return [
        (
            "1. Assunto",
            "Parecer técnico fundamentado na análise apresentada pelo Chat IA.",
        ),
        ("2. Análise", _clip(cleaned)),
        (
            "3. Conclusões e recomendações",
            "Adotar as recomendações do corpo técnico acima, sujeitas à "
            "validação do responsável técnico do empreendimento.",
        ),
        ("4. Limitações", _disclaimer()),
    ]


def _especificacao_sections(body: str) -> list[tuple[str, str]]:
    cleaned = _clean_md(body)
    return [
        (
            "1. Objeto",
            "Especificação técnica derivada da solução recomendada no Chat IA.",
        ),
        ("2. Requisitos técnicos", _clip(cleaned)),
        (
            "3. Critérios de aceitação",
            "Atender normas citadas, tolerâncias de projeto e inspeção "
            "pelo responsável técnico antes da liberação.",
        ),
        ("4. Responsabilidades", _disclaimer()),
    ]


def _checklist_sections(body: str) -> list[tuple[str, str]]:
    cleaned = _clean_md(body)
    return [
        (
            "1. Finalidade",
            "Checklist de verificação técnica gerado a partir da resposta do Chat IA.",
        ),
        ("2. Itens e critérios", _clip(cleaned)),
        (
            "3. Registro",
            "Marcar conformidade / não conformidade em campo e anexar evidências "
            "(fotos, medições, documentos).",
        ),
        ("4. Responsabilidades", _disclaimer()),
    ]


def _sectioned_content(kind: DocKind, body: str) -> list[tuple[str, str]]:
    if kind == "trd":
        return _trd_sections(body)
    if kind == "memorial":
        return _memorial_sections(body)
    if kind == "parecer":
        return _parecer_sections(body)
    if kind == "especificacao":
        return _especificacao_sections(body)
    if kind == "checklist":
        return _checklist_sections(body)
    return _trd_sections(body)


def _doc_meta(
    *,
    kind: DocKind,
    title: str | None,
    discipline: str | None,
) -> dict[str, Any]:
    heading = title or _TITLE.get(kind, _TITLE["resposta"])
    code = _DOC_CODE.get(kind, "DOC")
    stamp = datetime.now().strftime("%Y%m%d")
    return {
        "titulo": heading,
        "numero_laudo": f"{code}-{stamp}",
        "objeto": heading,
        "data_vistoria": datetime.now().strftime("%d/%m/%Y"),
        "disciplina": discipline or "—",
    }


def build_chat_docx(
    *,
    kind: DocKind,
    text: str,
    title: str | None = None,
    discipline: str | None = None,
    source_question: str | None = None,
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    from core.inspection_report.docx_export import (
        BLACK,
        BLUE,
        GRAY,
        _add_heading,
        _add_paragraph,
        _enable_update_fields_on_open,
        _set_run_font,
        _setup_header_footer,
    )
    from core.inspection_report.format_utils import format_generated_at

    generated_at = format_generated_at(datetime.now())
    meta = _doc_meta(kind=kind, title=title, discipline=discipline)
    doc = Document()
    _setup_header_footer(doc, content=meta, generated_at=generated_at)
    _enable_update_fields_on_open(doc)

    heading = meta["titulo"]
    _add_heading(doc, heading, level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Documento nº {meta['numero_laudo']}"
        + (f" · Disciplina: {discipline}" if discipline else "")
        + f" · Gerado em {generated_at}"
    )
    _set_run_font(run, size=9, color=GRAY)

    _add_paragraph(
        doc,
        "Layout institucional alinhado ao módulo de Laudos (cabeçalho, rodapé, "
        "marca d'água e tipografia Times New Roman).",
        first_indent=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9,
        space_after=12,
        color=GRAY,
    )

    if source_question:
        _add_heading(doc, "1. Solicitação", level=1)
        _add_paragraph(doc, source_question.strip(), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    if kind in _SECTIONED:
        base_n = 2 if source_question else 1
        for i, (sec_title, sec_body) in enumerate(_sectioned_content(kind, text), start=base_n):
            # sec_title já vem numerado ("1. Objeto") — renumerar de forma simples
            label = re.sub(r"^\d+\.\s*", "", sec_title)
            _add_heading(doc, f"{i}. {label}", level=1)
            for para in sec_body.split("\n"):
                if para.strip():
                    _add_paragraph(doc, para.strip(), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    else:
        _add_heading(doc, "2. Conteúdo técnico" if source_question else "1. Conteúdo técnico", level=1)
        for bkind, content in _split_blocks(text):
            if bkind == "heading":
                _add_heading(doc, content, level=2)
            elif bkind == "table_row":
                p = doc.add_paragraph(content)
                for r in p.runs:
                    _set_run_font(r, size=9, color=BLACK)
                    r.font.name = "Consolas"
            else:
                for para in content.split("\n\n"):
                    if para.strip():
                        _add_paragraph(doc, para.strip(), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    _add_paragraph(
        doc,
        "Documento gerado automaticamente pelo IA Server Santos. "
        "Validar com profissional habilitado antes do uso oficial.",
        first_indent=False,
        size=8,
        space_after=4,
        color=GRAY,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_chat_pdf(
    *,
    kind: DocKind,
    text: str,
    title: str | None = None,
    discipline: str | None = None,
    source_question: str | None = None,
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    from core.inspection_report.analytics import prepare_watermark_png
    from core.inspection_report.format_utils import (
        format_generated_at,
        header_meta_lines,
    )
    from core.inspection_report.pdf_export import _esc, _make_page_callbacks, _styles
    from core.system.company_profile import (
        get_company_profile,
        load_company_brasao,
        load_company_logo,
    )

    generated_at = format_generated_at(datetime.now())
    meta = _doc_meta(kind=kind, title=title, discipline=discipline)
    company = get_company_profile()
    org = company.display_name() or company.razao_social or "Empresa responsável"
    company_line = " | ".join(
        x for x in (company.endereco_linha(), company.site or company.email) if x
    ) or org
    footer_ref = str(meta.get("numero_laudo") or meta.get("titulo") or "Documento técnico")
    meta_lines = header_meta_lines(meta, generated_at=generated_at)

    logo_bytes = load_company_logo()
    brasao_raw = load_company_brasao()
    watermark_bytes = None
    if brasao_raw:
        watermark_bytes = prepare_watermark_png(
            brasao_raw, size_px=1800, opacity=0.06, max_width_px=1800
        ) or brasao_raw

    styles = _styles()
    story: list[Any] = []
    story.append(Paragraph(_esc(meta["titulo"]), styles["h1"]))
    story.append(
        Paragraph(
            _esc(
                f"Documento nº {meta['numero_laudo']}"
                + (f" · Disciplina: {discipline}" if discipline else "")
                + f" · Gerado em {generated_at}"
            ),
            styles["meta"],
        )
    )
    story.append(
        Paragraph(
            _esc(
                "Layout institucional alinhado ao módulo de Laudos "
                "(cabeçalho, rodapé, marca d'água e tipografia)."
            ),
            styles["meta"],
        )
    )
    story.append(Spacer(1, 8))

    if source_question:
        story.append(Paragraph(_esc("1. Solicitação"), styles["h1"]))
        story.append(Paragraph(_esc(source_question.strip()), styles["body"]))

    if kind in _SECTIONED:
        base_n = 2 if source_question else 1
        for i, (sec_title, sec_body) in enumerate(_sectioned_content(kind, text), start=base_n):
            label = re.sub(r"^\d+\.\s*", "", sec_title)
            story.append(Paragraph(_esc(f"{i}. {label}"), styles["h1"]))
            story.append(Paragraph(_esc(sec_body), styles["body"]))
    else:
        story.append(
            Paragraph(
                _esc("2. Conteúdo técnico" if source_question else "1. Conteúdo técnico"),
                styles["h1"],
            )
        )
        for bkind, content in _split_blocks(text):
            if bkind == "heading":
                story.append(Paragraph(_esc(content), styles["h2"]))
            elif bkind == "table_row":
                story.append(Paragraph(_esc(content), styles["meta_left"]))
            else:
                for para in content.split("\n\n"):
                    if para.strip():
                        story.append(Paragraph(_esc(para.strip()), styles["body"]))

    story.append(Spacer(1, 10))
    story.append(
        Paragraph(
            _esc(
                "Documento gerado automaticamente pelo IA Server Santos. "
                "Validar com profissional habilitado antes do uso oficial."
            ),
            styles["meta"],
        )
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.0 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.3 * cm,
    )
    on_page = _make_page_callbacks(
        org=org,
        company_line=company_line,
        footer_ref=footer_ref,
        logo_bytes=logo_bytes,
        meta_lines=meta_lines,
        watermark_bytes=watermark_bytes or brasao_raw,
    )
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return buf.getvalue()


def build_chat_document(
    *,
    kind: DocKind,
    fmt: DocFormat,
    text: str,
    title: str | None = None,
    discipline: str | None = None,
    source_question: str | None = None,
) -> tuple[bytes, str, str]:
    safe_kind = kind if kind in _TITLE else "resposta"
    if fmt == "docx":
        data = build_chat_docx(
            kind=safe_kind,  # type: ignore[arg-type]
            text=text,
            title=title,
            discipline=discipline,
            source_question=source_question,
        )
        filename = f"{safe_kind}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        return (
            data,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename,
        )
    data = build_chat_pdf(
        kind=safe_kind,  # type: ignore[arg-type]
        text=text,
        title=title,
        discipline=discipline,
        source_question=source_question,
    )
    filename = f"{safe_kind}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    return data, "application/pdf", filename
